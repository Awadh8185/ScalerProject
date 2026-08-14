#!/usr/bin/env python3
"""Redact PII in a DOCX while retaining its Word layout and structure."""
from __future__ import annotations

import argparse
import csv
import hashlib
import re
from collections import Counter
from pathlib import Path
from typing import Callable

from docx import Document


EMAIL = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
# 10–15 digits, allowing normal international separators. The digit-count
# look-ahead prevents SSNs and short reference numbers from being treated as phones.
PHONE = re.compile(r"(?<!\w)(?=(?:[^0-9;\r\n]*\d){10,15}(?!\d))(?:\+\s?)?[\d(][\d().\s-]{8,24}\d(?!\w)")
SSN = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
IP = re.compile(r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|1?\d?\d)\b")
CARD = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")
# These include their labels because Python's regex engine does not allow a
# variable-length look-behind (for example, "DOB" versus "date of birth").
DOB = re.compile(r"(?i)\b(?:dob|date of birth)\s*[:=-]\s*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}")
ADDRESS = re.compile(r"(?i)\b(?:address|mailing address|residential address)\s*[:=-]\s*[^;\n]+")
LABELED_NAME = re.compile(r"(?i)\b(?:name|customer|client|applicant|employee)\s*[:=-]\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}")
LABELED_COMPANY = re.compile(r"(?i)(?<!\[)\b(?:company|employer|organisation|organization)\s*[:=-]\s*[A-Za-z][^;\n]+")
# Extra high-precision heuristics for values that are not introduced by a label.
# They avoid treating ordinary capitalised prose as a person's or company's name.
HONORIFIC_NAME = re.compile(r"\b(?:Mr|Ms|Mrs|Dr)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}")
COMPANY_LEGAL = re.compile(
    r"\b(?:[A-Z][A-Za-z0-9&'.,-]*\s+){0,6}(?:Private\s+Limited|Public\s+Limited|Pvt\.?\s+Ltd\.?|Limited|Ltd\.?|LLP|Inc\.?|Corporation|Corp\.?)\b"
)
CONTACT_PERSON = re.compile(
    r"(?i)(\bcontact person\s*:\s*)([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3}(?:\s*/\s*[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})*)"
)
INLINE_ADDRESS = re.compile(
    r"(?i)(?<!\w)\d{1,4},?\s+[^;\n]{10,150}?(?=\s+(?:telephone|tel\.?|email|e-mail|website|contact person)\s*:)"
)

FIRST = ("Aarav", "Diya", "Kabir", "Meera", "Arjun", "Isha", "Rohan", "Anaya")
LAST = ("Sharma", "Patel", "Singh", "Gupta", "Kumar", "Das", "Mehta", "Roy")
COMPANIES = ("Northstar Solutions Pvt. Ltd.", "Summit Data Systems Ltd.", "Bluewave Services Pvt. Ltd.")


def stable_choice(values: tuple[str, ...], original: str) -> str:
    return values[int(hashlib.sha256(original.encode()).hexdigest(), 16) % len(values)]


def fake(kind: str, original: str) -> str:
    key = int(hashlib.sha256(original.encode()).hexdigest()[:8], 16)
    if kind == "EMAIL":
        return f"user{key % 100000}@example.com"
    if kind == "PHONE":
        return f"+91 900{key % 10000000:07d}"
    if kind == "SSN":
        return f"{100 + key % 800:03d}-{10 + key % 80:02d}-{1000 + key % 8000:04d}"
    if kind == "IP":
        return f"203.0.113.{1 + key % 254}"  # TEST-NET-3; never a real host
    if kind == "CARD":
        body = "400000" + f"{key % 1_000_000_000:09d}"
        check_digit = next(str(digit) for digit in range(10) if luhn(body + str(digit)))
        value = body + check_digit
        return " ".join(value[index:index + 4] for index in range(0, 16, 4))
    if kind == "DOB":
        return f"{1 + key % 28:02d}/{1 + (key // 29) % 12:02d}/1990"
    if kind == "ADDRESS":
        return f"{1 + key % 999} Example Road, Sample City, 400001"
    if kind == "NAME":
        return f"{stable_choice(FIRST, original)} {stable_choice(LAST, original[::-1])}"
    if kind == "COMPANY":
        return stable_choice(COMPANIES, original)
    raise ValueError(kind)


def luhn(value: str) -> bool:
    digits = [int(c) for c in re.sub(r"\D", "", value)]
    if not 13 <= len(digits) <= 19:
        return False
    total = 0
    for index, digit in enumerate(reversed(digits)):
        if index % 2:
            digit *= 2
            digit -= 9 if digit > 9 else 0
        total += digit
    return total % 10 == 0


def redact_text(text: str, counts: Counter) -> str:
    # Contact names commonly appear in prospectuses as a list separated by '/'.
    def contact_replacement(match: re.Match) -> str:
        names = re.split(r"\s*/\s*", match.group(2))
        replacements = []
        for name in names:
            counts["NAME"] += 1
            replacements.append(fake("NAME", name))
        return match.group(1) + "[NAME: " + " / ".join(replacements) + "]"

    text = CONTACT_PERSON.sub(contact_replacement, text)
    rules: list[tuple[str, re.Pattern, Callable[[str], bool] | None]] = [
        ("EMAIL", EMAIL, None), ("SSN", SSN, None), ("IP", IP, None),
        # The eight-digits/two-digits form is a prospectus reference number, not a phone.
        ("CARD", CARD, luhn), ("PHONE", PHONE, lambda value: not luhn(value) and not re.fullmatch(r"\d{8}-\d{2}", value)), ("DOB", DOB, None),
        ("ADDRESS", ADDRESS, None), ("ADDRESS", INLINE_ADDRESS, None),
        ("COMPANY", COMPANY_LEGAL, None), ("COMPANY", LABELED_COMPANY, None),
        ("NAME", HONORIFIC_NAME, None), ("NAME", LABELED_NAME, None),
    ]
    for kind, pattern, predicate in rules:
        def replacement(match: re.Match, kind: str = kind) -> str:
            value = match.group(0)
            if predicate and not predicate(value):
                return value
            counts[kind] += 1
            return f"[{kind}: {fake(kind, value)}]"
        text = pattern.sub(replacement, text)
    return text


def paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def redact_docx(source: Path, destination: Path) -> Counter:
    document = Document(source)
    counts: Counter = Counter()
    for paragraph in paragraphs(document):
        # Replacing paragraph.text preserves content but normalises character-level formatting.
        redacted = redact_text(paragraph.text, counts)
        if redacted != paragraph.text:
            paragraph.text = redacted
    document.save(destination)
    return counts


def evaluate(gold_csv: Path) -> dict[str, float]:
    """Evaluate rows with columns: actual,predicted (1 is PII/redacted)."""
    matrix = Counter()
    with gold_csv.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            actual, predicted = int(row["actual"]), int(row["predicted"])
            matrix[(actual, predicted)] += 1
    tp, fp, fn, tn = matrix[(1, 1)], matrix[(0, 1)], matrix[(1, 0)], matrix[(0, 0)]
    precision = tp / (tp + fp) if tp + fp else 0.0
    recall = tp / (tp + fn) if tp + fn else 0.0
    accuracy = (tp + tn) / sum(matrix.values()) if matrix else 0.0
    return {"tp": tp, "fp": fp, "fn": fn, "tn": tn, "accuracy": accuracy, "precision": precision, "recall": recall}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path, nargs="?", help="source .docx")
    parser.add_argument("output", type=Path, nargs="?", help="redacted .docx")
    parser.add_argument("--evaluate", type=Path, metavar="CSV", help="evaluate a labelled CSV instead")
    args = parser.parse_args()
    if args.evaluate:
        for name, value in evaluate(args.evaluate).items():
            print(f"{name}: {value:.4f}" if isinstance(value, float) else f"{name}: {value}")
    elif args.input and args.output:
        if args.input.suffix.lower() != ".docx":
            parser.error("input must be a .docx file")
        counts = redact_docx(args.input, args.output)
        print(f"Saved {args.output}")
        print("Redactions: " + ", ".join(f"{key}={value}" for key, value in sorted(counts.items())))
    else:
        parser.error("provide input and output, or --evaluate CSV")


if __name__ == "__main__":
    main()
